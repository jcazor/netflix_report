import asyncio
import uuid
import docx
import smtplib, os
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.utils import COMMASPACE, formatdate
from email import encoders
from komlogd.api import logging
from komlogd.api.transfer_methods import transfermethod
from komlogd.api.protocol.model.types import Datapoint, Datasource

# Commands definition, by default we set sysstat commands count to 3
UPTIME = 'uptime'
DMESG = 'dmesg | tail'
VMSTAT = 'vmstat 1 3'
MPSTAT = 'mpstat -P ALL 1 3'
PIDSTAT = 'pidstat 1 3'
IOSTAT = 'iostat -xz 1 3'
FREE = 'free -m'
SAR_DEV = 'sar -n DEV 1 3'
SAR_TCP = 'sar -n TCP,ETCP 1 3'
TOP = 'top -b -n1'

# email vars
SERVER = 'smtp_server'
PORT = 587
USERNAME = 'username'
PASSWORD = 'pass'
FROM = 'me@domain.com'
TO =['email1@domain.com','....']


class Command:
    def __init__(self, command, stdout):
        self.command = command
        self.stdout = Datasource(uri=stdout)

class NetflixReport:
    def __init__(self):
        self.hostname = 'hostname' #you could use socket.gethostname() for example, to detect it
        BASE_URI = '.'.join(('netflix_checklist',self.hostname))
        self.alarm = Datapoint(uri='.'.join((BASE_URI,'_alarm')))
        self.commands = []
        self.commands.append(Command(UPTIME,'.'.join((BASE_URI,'uptime'))))
        self.commands.append(Command(DMESG,'.'.join((BASE_URI,'dmesg'))))
        self.commands.append(Command(VMSTAT,'.'.join((BASE_URI,'vmstat'))))
        self.commands.append(Command(MPSTAT,'.'.join((BASE_URI,'mpstat'))))
        self.commands.append(Command(PIDSTAT,'.'.join((BASE_URI,'pidstat'))))
        self.commands.append(Command(IOSTAT,'.'.join((BASE_URI,'iostat'))))
        self.commands.append(Command(FREE,'.'.join((BASE_URI,'free'))))
        self.commands.append(Command(SAR_DEV,'.'.join((BASE_URI,'sar_dev'))))
        self.commands.append(Command(SAR_TCP,'.'.join((BASE_URI,'sar_tcp'))))
        self.commands.append(Command(TOP,'.'.join((BASE_URI,'top'))))

report = NetflixReport()

@transfermethod(p_in={'alarm':report.alarm}, p_out={'commands':report.commands})
async def run_report(alarm, commands):
    ts = alarm.data.ix[-1]
    contents = {}
    for cmd in commands:
        try:
            p = await asyncio.create_subprocess_shell(cmd.command, stdout=asyncio.subprocess.PIPE, stderr = asyncio.subprocess.PIPE)
            output = await p.stdout.read()
        except Exception as e:
            logging.logger.error('Exception running command.')
            logging.logger.error(str(e))
        else:
            await p.wait()
            content = output.decode('utf-8')
            cmd.stdout.data[ts] = content
            contents[cmd.stdout.uri]=content
    doc = create_perf_report(contents)
    if doc:
        filename = '/tmp/'+uuid.uuid4().hex+'.docx'
        doc.save(filename)
        send_mail(send_from=FROM, send_to=TO, subject='Perf report', text='Perf report', files=[filename])
        os.remove(filename)

def create_perf_report(contents):
    ''' This function creates a basic .docx document. This code is for example purposes only,
        so we do not apply any text formatting, and the result is not very nice :) '''
    try:
        document=docx.Document()
        document.add_heading('Basic System Performance Report',0)
        for uri,content in contents.items():
            document.add_heading(uri,1)
            document.add_paragraph(content)
    except Exception as e:
        return None
    else:
        return document

def send_mail( send_from, send_to, subject, text, files=[]):
    msg = MIMEMultipart()
    msg['From'] = send_from
    msg['To'] = COMMASPACE.join(send_to)
    msg['Date'] = formatdate(localtime = True)
    msg['Subject'] = subject
    msg.attach( MIMEText(text) )
    for f in files:
        part = MIMEBase('application', "octet-stream")
        part.set_payload( open(f,"rb").read() )
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment; filename="{0}"'.format(os.path.basename(f)))
        msg.attach(part)
    smtp = smtplib.SMTP(SERVER,PORT)
    smtp.starttls()
    smtp.login(USERNAME,PASS)
    smtp.sendmail(send_from, send_to, msg.as_string())
    smtp.quit()

